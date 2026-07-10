"""
企业年报 PDF/文本解析模块
支持 PDF（多引擎 fallback）、TXT、DOCX 格式
"""
import os
import re
import io
from pathlib import Path
from typing import Optional, Dict, Any


class AnnualReportParser:
    """年报解析器 — 多格式、多引擎、健壮解析"""

    def __init__(self):
        self.supported_formats = [".pdf", ".txt", ".docx"]
        self.max_file_size_mb = 50

    def parse(
        self,
        file_path: str,
        file_bytes: Optional[bytes] = None,
    ) -> Optional[str]:
        """解析年报文件，返回纯文本内容"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.supported_formats:
            raise ValueError(
                f"不支持的文件格式：{suffix}，支持：{', '.join(self.supported_formats)}"
            )

        if file_bytes is not None:
            file_size_mb = len(file_bytes) / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                raise ValueError(
                    f"文件过大（{file_size_mb:.1f}MB），超过最大限制 {self.max_file_size_mb}MB"
                )

        if suffix == ".pdf":
            return self._parse_pdf(path, file_bytes)
        elif suffix == ".txt":
            return self._parse_txt(path, file_bytes)
        elif suffix == ".docx":
            return self._parse_docx(path, file_bytes)

        return None

    # ============================================================
    #  PDF 解析 — 多引擎 fallback
    # ============================================================
    def _parse_pdf(
        self, file_path: Path, file_bytes: Optional[bytes] = None
    ) -> Optional[str]:
        """PDF 多引擎 fallback：pdfplumber → PyPDF2 → pypdf → pdfminer"""
        engines = [
            ("pdfplumber", self._parse_pdf_pdfplumber),
            ("PyPDF2", self._parse_pdf_pypdf2),
            ("pypdf", self._parse_pdf_pypdf),
            ("pdfminer", self._parse_pdf_pdfminer),
        ]

        errors = []
        for engine_name, engine_func in engines:
            try:
                text = engine_func(file_path, file_bytes)
                if text and len(text.strip()) > 50:
                    return text
                elif text and len(text.strip()) > 0:
                    errors.append(f"{engine_name}: 提取内容过少")
            except ImportError:
                errors.append(f"{engine_name}: 未安装")
            except Exception as e:
                errors.append(f"{engine_name}: {e}")

        error_msg = "；".join(errors) if errors else "未知错误"
        raise RuntimeError(
            f"PDF文件无法解析。\n"
            f"已尝试引擎：{', '.join(e[0] for e in engines)}\n"
            f"错误详情：{error_msg}\n"
            f"建议：1) 确认PDF文件是否为扫描件（扫描件需OCR）；2) 安装 pdfplumber：pip install pdfplumber"
        )

    def _parse_pdf_pdfplumber(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        import pdfplumber
        source = io.BytesIO(file_bytes) if file_bytes else str(file_path)
        text = ""
        with pdfplumber.open(source) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception:
                    continue
        return text if text.strip() else None

    def _parse_pdf_pypdf2(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        from PyPDF2 import PdfReader
        source = io.BytesIO(file_bytes) if file_bytes else str(file_path)
        reader = PdfReader(source)
        text = ""
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            except Exception:
                continue
        return text if text.strip() else None

    def _parse_pdf_pypdf(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        from pypdf import PdfReader
        source = io.BytesIO(file_bytes) if file_bytes else str(file_path)
        reader = PdfReader(source)
        text = ""
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            except Exception:
                continue
        return text if text.strip() else None

    def _parse_pdf_pdfminer(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        from pdfminer.high_level import extract_text
        if file_bytes:
            text = extract_text(io.BytesIO(file_bytes))
        else:
            text = extract_text(str(file_path))
        return text if text and text.strip() else None

    # ============================================================
    #  TXT 解析
    # ============================================================
    def _parse_txt(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030", "utf-16", "latin-1"]
        raw = file_bytes if file_bytes else file_path.read_bytes()
        for enc in encodings:
            try:
                text = raw.decode(enc)
                return text
            except (UnicodeDecodeError, LookupError):
                continue
        raise RuntimeError(
            f"文本文件编码无法识别。已尝试编码：{', '.join(encodings)}\n"
            f"建议：将文件另存为 UTF-8 编码后重新上传"
        )

    # ============================================================
    #  DOCX 解析
    # ============================================================
    def _parse_docx(self, file_path: Path, file_bytes: Optional[bytes] = None) -> Optional[str]:
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("未安装 python-docx 库\n请运行：pip install python-docx")

        source = io.BytesIO(file_bytes) if file_bytes else str(file_path)
        doc = Document(source)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text += row_text + "\n"
        if not text.strip():
            raise RuntimeError("DOCX文件内容为空或无法提取")
        return text

    # ============================================================
    #  元数据提取
    # ============================================================
    def extract_metadata(self, file_path: str, file_bytes: Optional[bytes] = None) -> Dict[str, Any]:
        path = Path(file_path)
        file_size = 0
        if file_bytes:
            file_size = len(file_bytes)
        elif path.exists():
            file_size = path.stat().st_size
        return {
            "file_name": path.name,
            "file_size": file_size,
            "file_size_kb": round(file_size / 1024, 1),
            "file_size_mb": round(file_size / (1024 * 1024), 2),
            "file_format": path.suffix.lower().lstrip("."),
            "file_stem": path.stem,
        }

    # ============================================================
    #  辅助：从文件名推断企业信息
    # ============================================================
    def infer_company_info(self, file_name: str) -> Dict[str, str]:
        info = {"company_name": "", "year": ""}
        year_match = re.search(r"(20\d{2})", file_name)
        if year_match:
            info["year"] = year_match.group(1)
        clean_name = re.sub(r"20\d{2}.*$", "", file_name)
        clean_name = re.sub(
            r"(年度报告|年报|社会责任报告|ESG报告|环境报告|可持续发展报告"
            r"|年度|报告|股份有限公司|有限公司|集团|公司)",
            "",
            clean_name,
        )
        clean_name = clean_name.strip(" -_·")
        if clean_name:
            info["company_name"] = clean_name
        return info